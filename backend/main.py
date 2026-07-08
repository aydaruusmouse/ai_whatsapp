"""Telesom AI — chatbot with file + URL knowledge (Somali primary, English supported)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from dotenv import load_dotenv

ROOT = Path(__file__).resolve().parent
load_dotenv(ROOT / ".env", override=True)
load_dotenv(override=False)

import io
import os
import re
import shutil
import subprocess
import tempfile

import httpx
from bs4 import BeautifulSoup
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from openai import APIStatusError, OpenAI
from pydantic import BaseModel, Field

from flow_state import get_meta, reset_flow, set_meta
from knowledge import (
    add_knowledge,
    chunk_count,
    db_status,
    ensure_session,
    init_db,
    new_session,
    RetrievalResult,
    retrieve_context_result,
    total_chunk_count,
)
from vas_flow import handle_vas_flow
from vas_offers import text_mentions_known_offer

STATIC = ROOT / "static"

SYSTEM_PROMPT_SO = """Waxaad tahay kaaliyaha macaamiisha shirkadda Telesom (AI support bot).

HABKA JAWAABTA (raac si adag — ma jiro wax ka reeban):
1. Jawaab kasta waa inay KA TIMATO kaliya qaybta "Xogta hoose" (faylal la soo raray iyo links la geliyay).
2. Ka hor intaadan jawaabin, baadh oo akhri Dhammaan xogta hoose — waa ilahaaga kaliya ee runta ah.
3. Ka jawaab su'aasha KALIYA haddii jawaabtu si toos ah uga muuqato xogta hoose (qoraal ama macne toos ah).
4. Mamnuuc: ha u maleynin, ha ka sheegin wax aan xogta ku jirin, ha isticmaalin aqoontaada guud ee tababarka.
5. HADDII jawaabtu aanay si cad ugu jirin xogta hoose, soo celi EXACTLY kaliya: [HANDOFF]
   — ha qorin wax kale, ha bixin talo guud, ha sheegin wax aadan ka akhrin xogta.

Xeerar kale:
- Ha soo bandhigin liiska VAS offers (47 adeeg) — liiskaas waxaa bixiya nidaamka VAS marka la weydiiyo subscribe.
- Su'aalaha ku saabsan Telesom: soo koob, cad, ku saleysan xogta kaliya.

Luqadda: Af-Soomaali marka hore; Ingiriis haddii isticmaaluhu Ingiriis qoro.

Jawaabaha: gaaban, cad, ixtiraam leh — kaliya waxa xogta ku jira."""

HANDOFF_SENTINEL = "[HANDOFF]"

app = FastAPI(title="Telesom AI Chatbot")


@app.on_event("startup")
def _startup_db() -> None:
    init_db()


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_client: OpenAI | None = None
HANDOFF_WEBHOOK_URL = "https://support.telesom.com/index.php/api/webhook/whatsapp"
HANDOFF_REPLY = (
    "Macmiil, waxaan kugu xidheynaa qof ku caawiya fadlan noo sabir,\n"
    "Customer, we are connecting you with someone to assist you. Please be patient with us.\n\n"
    "Hadaad rabta inaa iska xidho wada hadalkan fadlan soo qor /end.\n"
    "If you want to end this conversation, please type /end."
)
HANDOFF_ACK = (
    "\n\nAsalaamu Calaykum,\n\n"
    "Ku soo dhawoow Xafiiska Daryeel Macaamiisha ee Telesom.\n"
    "Waad ku mahadsan tahay inaad nala soo xidhiidho. Codsigaagii waan helnay, "
    "waxaana laguugu soo jawaabi doonaa sida ugu dhakhso badan.\n\n"
    "Mahadsanid."
)
FULL_HANDOFF_REPLY = HANDOFF_REPLY + HANDOFF_ACK
HANDOFF_WAITING_REPLY = (
    "Macmiil, codsigaagii waan helnay — fadlan noo sabir, qof ayaa kuu caawin doona.\n"
    "We have received your message — please wait, someone will assist you shortly."
)
END_HANDOFF_REPLY = (
    "Wadahadalka waa la xidhay. Haddii aad wax kale u baahato Telesom, mar kale noo soo qor.\n"
    "Conversation ended. You can message us again anytime."
)

_TELESOM_SIGNALS = (
    "telesom", "zaad", "evc", "zaadplus", "fiber", "vas", "bundle", "bundles",
    "data", "internet", "sim", "recharge", "airtime", "lacag", "balance",
    "mobile", "wicitaan", "call", "sms", "4g", "5g", "broadband",
    "subscribe", "unsubscribe", "evcplus", "kaafiye", "mifi", "151",
    "cabasho", "complaint", "tixraac", "transaction", "sarif", "exchange",
    "shilling", "dollar", "xannib", "xayir", "block", "offer", "football",
    "live score", "newspaper", "ramadan", "evc plus", "zaad plus",
)

GREETING_REPLY = "kuso dhawaaw shirkada telesom maxaan kaa caawinaa macmiil"


def _mock_llm_enabled() -> bool:
    return os.getenv("MOCK_LLM", "").strip().lower() in ("1", "true", "yes", "on")


def _mock_reply(last_user: str, context: str, *, best_score: int = 0) -> str:
    """Local demo replies when OpenAI is not used (no API key billing)."""
    if not context.strip() or best_score <= 0:
        return HANDOFF_SENTINEL
    excerpt = context.strip()
    if len(excerpt) > 1200:
        excerpt = excerpt[:1200].rsplit(" ", 1)[0] + "…"
    return (
        "(Habka DEMO — MOCK_LLM; ma jiro OpenAI) Su'aasha: «"
        + last_user.strip()[:200]
        + "»\n\n"
        "Qayb ka mid ah xogta la geliyay:\n\n"
        + excerpt
    )


def _llm_requests_handoff(reply: str) -> bool:
    r = (reply or "").strip()
    return r == HANDOFF_SENTINEL or r.startswith(HANDOFF_SENTINEL)


def _knowledge_context_block(retrieval: RetrievalResult) -> str:
    context = retrieval.text
    if context.strip() and retrieval.best_score > 0:
        return (
            f"Xogta hoose — faylal & links la geliyay ({retrieval.chunks_included}/"
            f"{retrieval.total_chunks} qaybood oo la helay; "
            "ka jawaab KALIYA waxa halkan ku qoran — ha maleynin):\n\n"
            + context
        )
    return (
        "Xogta hoose: ma jiro qayb ku habboon su'aasha. "
        f"Soo celi {HANDOFF_SENTINEL} — ha qorin wax kale."
    )


def _reply_from_knowledge(
    sid: str,
    last_user: str,
    body: ChatBody,
    retrieval: RetrievalResult,
) -> dict[str, Any]:
    """Answer strictly from retrieved document/link chunks; handoff if LLM cannot cite knowledge."""
    if not retrieval.text.strip() or retrieval.best_score <= 0:
        return _handoff_response(sid, last_user, first_time=True)

    context_block = _knowledge_context_block(retrieval)
    api_messages = [
        {"role": "system", "content": SYSTEM_PROMPT_SO},
        {"role": "system", "content": context_block},
    ]
    for m in body.messages:
        if m.role in ("user", "assistant"):
            api_messages.append({"role": m.role, "content": m.content})

    if _mock_llm_enabled():
        choice = _mock_reply(last_user, retrieval.text, best_score=retrieval.best_score)
        if _llm_requests_handoff(choice):
            return _handoff_response(sid, last_user, first_time=True)
        return {
            "session_id": sid,
            "reply": choice,
            "chunks_in_memory": chunk_count(sid),
            "mock_llm": True,
            "knowledge_only": True,
        }

    default_model = (
        os.getenv("OPENAI_MODEL", "llama3.2")
        if _openai_base_url()
        else os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    )
    model = body.model or default_model
    try:
        resp = get_client().chat.completions.create(
            model=model,
            messages=api_messages,
            temperature=0,
        )
    except APIStatusError as e:
        msg = str(e)
        api_msg = ""
        if isinstance(e.body, dict):
            err = e.body.get("error")
            if isinstance(err, dict) and err.get("message"):
                api_msg = str(err["message"])
            elif isinstance(err, str):
                api_msg = err
        detail_tail = f" {api_msg}" if api_msg and api_msg not in msg else ""
        code = getattr(e, "status_code", None)
        if code == 401 or "401" in msg:
            raise HTTPException(
                status_code=401,
                detail=(
                    "Furaha API waa khalad (401). Hubi OPENAI_API_KEY + OPENAI_BASE_URL ee backend/.env "
                    "(OpenRouter: openrouter.ai/keys). Furaha chat-ka ku wadaagay waa in la cusboonaysiiyo."
                    + detail_tail
                ),
            ) from e
        if code == 402 or "402" in msg or "credits" in msg.lower():
            raise HTTPException(
                status_code=402,
                detail=(
                    "Lacag / credits OpenRouter ma filna (402). Ku dar credits ama dooro moodal bilaash "
                    "(openrouter.ai/models)." + detail_tail
                ),
            ) from e
        if code == 429 or "429" in msg or "rate" in msg.lower():
            raise HTTPException(
                status_code=429,
                detail="Xadka codsiyada / rate limit. Sug daqiiqad ama hubi xadka OpenRouter." + detail_tail,
            ) from e
        raise HTTPException(
            status_code=502,
            detail=f"API qalad ({code or '?'}): {msg}{detail_tail}",
        ) from e
    except Exception as e:
        msg = str(e)
        if "401" in msg or "invalid_api_key" in msg or "Incorrect API key" in msg or "Unauthorized" in msg:
            raise HTTPException(
                status_code=401,
                detail=(
                    "Furaha API waa khalad ama la diiday (401). Hubi backend/.env — OpenRouter: openrouter.ai/keys"
                ),
            ) from e
        raise HTTPException(502, f"OpenAI qalad: {msg}") from e

    choice = resp.choices[0].message.content or ""
    if _llm_requests_handoff(choice):
        return _handoff_response(sid, last_user, first_time=True)
    return {
        "session_id": sid,
        "reply": choice,
        "chunks_in_memory": chunk_count(sid),
        "mock_llm": False,
        "knowledge_only": True,
    }


def _normalize_text(s: str) -> str:
    return " ".join((s or "").strip().lower().split())


def _is_end_handoff(text: str) -> bool:
    q = _normalize_text(text)
    return q in {"/end", "end", "/end."}


def _is_service_inquiry(text: str) -> bool:
    """Su'aal ku saabsan adeeg gaar ah oo Telesom ah."""
    q = _normalize_text(text)
    if any(
        w in q
        for w in (
            "adeeg", "adeega", "adeegyada", "muxuu", "maxaa", "qabtaa", "qabato",
            "qabataa", "bixiyaa", "bixiyo", "service", "subscribe", "unsubscribe",
        )
    ):
        return True
    return text_mentions_known_offer(text)


def _is_telesom_related(text: str) -> bool:
    q = _normalize_text(text)
    if len(q) < 2:
        return True
    if _is_greeting(text) or _should_handoff_to_agent(text):
        return True
    if _is_general_telesom_services_query(text):
        return True
    if _is_service_inquiry(text):
        return True
    if any(s in q for s in _TELESOM_SIGNALS):
        return True
    if re.search(r"\d{6,}", q):
        return True
    return False


def _handoff_response(session_id: str, user_message: str, *, first_time: bool) -> dict[str, Any]:
    _notify_handoff_webhook(session_id, user_message)
    if first_time:
        set_meta(session_id, flow_state="agent_handoff", flow_data={"reason": "handoff"})
        reply = FULL_HANDOFF_REPLY
    else:
        reply = HANDOFF_WAITING_REPLY
    return {
        "session_id": session_id,
        "reply": reply,
        "chunks_in_memory": chunk_count(session_id),
        "mock_llm": False,
        "handoff_to_agent": True,
        "handoff_webhook_url": HANDOFF_WEBHOOK_URL,
    }


def _should_handoff_to_agent(text: str) -> bool:
    q = _normalize_text(text)
    triggers = (
        "connect with agent",
        "connect me with agent",
        "connect with me agent",
        "connect agent",
        "agent",
        "igu xidh kaaliye",
        "igu xidh caawiye",
        "ii xidh kaaliye",
        "ii xidh caawiye",
    )
    if any(t in q for t in triggers):
        return True
    # Flexible English intent: message includes both words anywhere.
    if "connect" in q and "agent" in q:
        return True
    # Flexible Somali intent.
    if ("xidh" in q or "xir" in q) and ("kaaliye" in q or "caawiye" in q):
        return True
    return False


def _is_greeting(text: str) -> bool:
    q = _normalize_text(text)
    return q in {
        "hi",
        "hello",
        "asc",
        "as",
        "asalamu calaykum",
        "assalamu calaykum",
        "salaam",
    }


def _is_general_telesom_services_query(text: str) -> bool:
    q = _normalize_text(text)
    so_hits = ("telesom" in q and ("adeeg" in q or "adeegyada" in q))
    en_hits = ("telesom" in q and ("service" in q or "services" in q))
    broad_intent = any(
        p in q
        for p in (
            "iga caawi adeegyada telesom",
            "adeegyada telesom",
            "telesom services",
            "help me with telesom services",
            "maxay telesom qabataa",
            "maxay telesom qabato",
            "waxa telesom qabato",
            "waxa telesom sameyso",
            "waxa ay telesom qabato",
            "what does telesom",
            "what is telesom",
        )
    )
    return so_hits or en_hits or broad_intent


def _notify_handoff_webhook(session_id: str, user_message: str) -> None:
    """Forward user handoff request to Chatwoot webhook."""
    twilio_account_sid = (os.getenv("TWILIO_ACCOUNT_SID") or "").strip()
    twilio_from = (os.getenv("TWILIO_WHATSAPP_FROM") or "").strip()
    twilio_auth_token = (os.getenv("TWILIO_AUTH_TOKEN") or "").strip()
    payload = {
        "session_id": session_id,
        "event": "agent_handoff_requested",
        "message": user_message,
        # Twilio-like fields for functions expecting inbound WhatsApp webhook shape.
        "AccountSid": twilio_account_sid,
        "From": twilio_from,
        "Body": user_message,
        # Keep auth token optional and only if explicitly configured.
        # NOTE: avoid forwarding secrets unless your receiver strictly requires it.
        "AuthToken": twilio_auth_token if twilio_auth_token else None,
    }
    payload = {k: v for k, v in payload.items() if v is not None}
    try:
        r = httpx.post(HANDOFF_WEBHOOK_URL, json=payload, timeout=12.0)
        r.raise_for_status()
    except Exception:
        # Do not block user handoff message if webhook temporarily fails.
        pass


def _openai_base_url() -> str:
    """Sida http://127.0.0.1:11434/v1 (Ollama) ama LM Studio OpenAI-compatible endpoint."""
    u = (os.getenv("OPENAI_BASE_URL") or "").strip().rstrip("/")
    return u


def _is_local_llm_base(base: str) -> bool:
    """Ollama / LM Studio oo isku xiran kumbuyuutarka — furaha API waa madhan / dummy."""
    b = base.lower()
    return (
        "localhost" in b
        or "127.0.0.1" in b
        or "0.0.0.0" in b
        or b.startswith("http://192.168.")
        or "host.docker.internal" in b
    )


def _is_placeholder_openai_key(key: str) -> bool:
    """Furaha runta ah waa dherer dheer (sk-... ama sk-proj-... tusaale ahaan waa gaaban)."""
    k = (key or "").strip()
    if not k:
        return True
    lower = k.lower()
    if lower in (
        "sk-...",
        "sk-proj-...",
        "sk-....",
        "your-key-here",
        "replace_me",
        "changeme",
    ):
        return True
    if "..." in k and len(k) < 30:
        return True
    # sk-or-v1-... (OpenRouter) iyo furayaasha dhererka leh waa sax
    if k.startswith("sk-") and not k.startswith("sk-or-v1-") and len(k) < 40:
        return True
    return False


def get_client() -> OpenAI:
    global _client
    if _client is None:
        base = _openai_base_url()
        key = (os.getenv("OPENAI_API_KEY") or "").strip()
        if base:
            if not key or _is_placeholder_openai_key(key):
                if _is_local_llm_base(base):
                    key = "ollama"
                else:
                    raise HTTPException(
                        status_code=503,
                        detail=(
                            "OPENAI_API_KEY waa madhan ama tusaale, laakiin OPENAI_BASE_URL waa adeeg "
                            "internet (Together, OpenRouter, iwm). Ku dheji furaha dhabta ah ee adeegga "
                            "ee .env — tusaale Together: api.together.xyz, OpenRouter: openrouter.ai/keys."
                        ),
                    )
            kwargs: dict = {"api_key": key, "base_url": base}
            if "openrouter.ai" in base:
                kwargs["default_headers"] = {
                    "HTTP-Referer": os.getenv(
                        "OPENROUTER_HTTP_REFERER",
                        "http://localhost:8000",
                    ),
                    "X-Title": os.getenv("OPENROUTER_APP_NAME", "Telesom AI"),
                }
            _client = OpenAI(**kwargs)
        else:
            if not key:
                raise HTTPException(
                    status_code=503,
                    detail=(
                        "OPENAI_API_KEY ma jiro. Faylka: backend/.env — ku qor furaha dhabta ah "
                        "https://platform.openai.com/api-keys — ama OPENAI_BASE_URL (Ollama) — ama MOCK_LLM=true."
                    ),
                )
            if _is_placeholder_openai_key(key):
                raise HTTPException(
                    status_code=503,
                    detail=(
                        "OPENAI_API_KEY wali waa tusaale ama gaaban yahay. Samee hal ka mid ah: "
                        "(1) Ku dheji furaha dhabta ah https://platform.openai.com/api-keys — "
                        "(2) Ama Ollama: OPENAI_BASE_URL=http://127.0.0.1:11434/v1 iyo OPENAI_MODEL=llama3.2 — "
                        "(3) Ama MOCK_LLM=true. Hubi in .env uu ku yaalo backend/.env"
                    ),
                )
            _client = OpenAI(api_key=key)
    return _client


class SessionBody(BaseModel):
    session_id: str | None = None
    customer_phone: str | None = None


class UrlBody(SessionBody):
    url: str = Field(..., min_length=4)


class ChatMessage(BaseModel):
    role: str
    content: str


class ChatBody(SessionBody):
    messages: list[ChatMessage]
    model: str | None = "gpt-4o-mini"


@app.post("/api/session")
def create_session(body: SessionBody | None = None):
    sid = new_session()
    if body and body.customer_phone:
        set_meta(sid, customer_phone=body.customer_phone.strip())
    return {"session_id": sid}


@app.post("/api/session/clear")
def api_clear_session(body: SessionBody):
    """Clear chat + VAS flow for this session; keep uploaded documents/links."""
    sid = ensure_session(body.session_id)
    reset_flow(sid)
    return {"session_id": sid, "chunks": total_chunk_count()}


@app.get("/api/knowledge/status")
def knowledge_status():
    """Global knowledge pool — not tied to browser session."""
    return {"chunks": total_chunk_count()}


@app.get("/api/session/status")
def session_status(session_id: str | None = None):
    sid = ensure_session(session_id)
    return {"session_id": sid, "chunks": total_chunk_count()}


def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _read_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_doc(data: bytes) -> str:
    """Legacy Word .doc — macOS textutil or antiword if installed."""
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        path = tmp.name
    try:
        if shutil.which("textutil"):
            proc = subprocess.run(
                ["textutil", "-stdout", "-convert", "txt", path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                return proc.stdout
        if shutil.which("antiword"):
            proc = subprocess.run(
                ["antiword", path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                return proc.stdout
        raise ValueError("legacy .doc")
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


_SUPPORTED_UPLOAD_SUFFIXES = (".pdf", ".txt", ".md", ".csv", ".doc", ".docx")


@app.post("/api/upload")
async def upload_file(
    file: UploadFile = File(...),
    session_id: str | None = Form(None),
):
    sid = ensure_session(session_id)
    name = file.filename or "upload"
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "Faylka waa madhan.")

    lower = name.lower()
    if not lower.endswith(_SUPPORTED_UPLOAD_SUFFIXES):
        raise HTTPException(
            400,
            "Nooca faylka lama taageero. Isticmaal PDF, TXT, MD, CSV, DOC, ama DOCX.",
        )

    try:
        if lower.endswith(".pdf"):
            text = _read_pdf(raw)
        elif lower.endswith(".docx"):
            text = _read_docx(raw)
        elif lower.endswith(".doc"):
            text = _read_doc(raw)
        else:
            text = raw.decode("utf-8", errors="replace")
    except ValueError as e:
        raise HTTPException(
            400,
            "Faylka DOC (Word hore) lama akhrin karo. Ku badel DOCX ama PDF oo mar kale soo rar.",
        ) from e
    except Exception as e:
        raise HTTPException(
            400,
            f"Faylka lama akhrin karo: {e!s}",
        ) from e

    if not text.strip():
        raise HTTPException(
            400,
            "Faylka waa madhan ama qoraal lama helin. Isticmaal PDF, TXT, MD, CSV, DOC, ama DOCX.",
        )

    n = add_knowledge(text, source=name)
    return {"session_id": sid, "chunks_added": n, "filename": name}


@app.post("/api/ingest-url")
async def ingest_url(body: UrlBody):
    sid = ensure_session(body.session_id)
    url = body.url.strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(400, "URL waa inuu ku bilaabmaa http:// ama https://")

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
            r = await client.get(url, headers={"User-Agent": "TelesomAI-Bot/1.0"})
            r.raise_for_status()
            ctype = (r.headers.get("content-type") or "").lower()
            if "pdf" in ctype:
                text = _read_pdf(r.content)
            else:
                html = r.text
                soup = BeautifulSoup(html, "html.parser")
                for tag in soup(["script", "style", "noscript"]):
                    tag.decompose()
                text = soup.get_text(separator="\n")
    except httpx.HTTPError as e:
        raise HTTPException(502, f"URL lama soo gelin karo: {e!s}") from e

    if not text.strip():
        raise HTTPException(422, "Bogga lama akhrin karin (wax qoraal ah ma jirin).")

    n = add_knowledge(text, source=url)
    return {"session_id": sid, "chunks_added": n, "url": url}


@app.post("/api/chat")
def chat(body: ChatBody):
    sid = ensure_session(body.session_id)
    if not body.messages:
        raise HTTPException(400, "Farriimaha waa la rabaa.")

    last_user = next(
        (m.content for m in reversed(body.messages) if m.role == "user"),
        "",
    )

    meta = get_meta(sid)
    if meta["flow_state"] == "agent_handoff":
        if _is_end_handoff(last_user):
            reset_flow(sid)
            return {
                "session_id": sid,
                "reply": END_HANDOFF_REPLY,
                "chunks_in_memory": chunk_count(sid),
                "mock_llm": False,
                "handoff_ended": True,
            }
        if not _is_telesom_related(last_user):
            return _handoff_response(sid, last_user, first_time=False)
        reset_flow(sid)

    if _should_handoff_to_agent(last_user):
        return _handoff_response(sid, last_user, first_time=True)

    # Mid-flow API states (VAS subscribe, block txn, fiber) take priority.
    meta = get_meta(sid)
    in_api_flow = meta["flow_state"] != "idle"
    if in_api_flow:
        vas_reply = handle_vas_flow(
            sid,
            last_user,
            body.customer_phone,
            skip_vas_menu=_is_general_telesom_services_query(last_user) or _is_service_inquiry(last_user),
        )
        if vas_reply is not None:
            return {
                "session_id": sid,
                "reply": vas_reply,
                "chunks_in_memory": chunk_count(sid),
                "mock_llm": False,
                "vas_flow": True,
            }

    if _is_greeting(last_user):
        return {
            "session_id": sid,
            "reply": GREETING_REPLY,
            "chunks_in_memory": chunk_count(sid),
            "mock_llm": False,
            "greeting": True,
        }

    # API intents (exchange, block wrong transaction, fiber, VAS menu)
    # should run before document knowledge replies.
    vas_reply = handle_vas_flow(
        sid,
        last_user,
        body.customer_phone,
        skip_vas_menu=_is_general_telesom_services_query(last_user) or _is_service_inquiry(last_user),
    )
    if vas_reply is not None:
        return {
            "session_id": sid,
            "reply": vas_reply,
            "chunks_in_memory": chunk_count(sid),
            "mock_llm": False,
            "vas_flow": True,
        }

    n_chunks = total_chunk_count()
    if n_chunks > 0:
        retrieval = retrieve_context_result(sid, last_user)
        if retrieval.best_score > 0 and retrieval.text.strip():
            return _reply_from_knowledge(sid, last_user, body, retrieval)

    return _handoff_response(sid, last_user, first_time=True)


if STATIC.exists():
    app.mount("/assets", StaticFiles(directory=str(STATIC)), name="assets")


@app.get("/")
def index():
    index_path = STATIC / "index.html"
    if index_path.is_file():
        return FileResponse(index_path)
    return {"message": "Ku dar static/index.html — ama u isticmaal /docs"}


@app.get("/health")
def health():
    base = bool(_openai_base_url())
    key = bool((os.getenv("OPENAI_API_KEY") or "").strip())
    if _mock_llm_enabled():
        backend = "mock"
    elif base:
        backend = "openai_compatible"
    else:
        backend = "openai"
    db = db_status()
    return {
        "ok": True,
        "mock_llm": _mock_llm_enabled(),
        "llm_backend": backend,
        "openai_compatible_url_set": base,
        "openai_api_key_set": key,
        **db,
    }
